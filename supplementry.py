#!/usr/bin/env python
# coding: utf-8

from bs4 import BeautifulSoup
from ebooklib import epub
import ebooklib
import docx
from difflib import SequenceMatcher
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import resolve1
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdftypes import PDFNotImplementedError
from io import StringIO
from nltk.corpus import stopwords
from nltk.cluster.util import cosine_distance
import numpy as np
import networkx as nx
from rake_nltk import Rake
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words
rake_nltk_var = Rake()
language = "english"
sentence_count = 10


def get_summary(text, top_n=5):
    summarize_text = ""
    parser = PlaintextParser(text, Tokenizer(language))
    summarizer_2 = LsaSummarizer(Stemmer(language))
    summarizer_2.stop_words = get_stop_words(language)
    summary_2 = summarizer_2(parser.document, sentence_count)
    for sentence in summary_2:
        summarize_text += f"{sentence}"

    rake_nltk_var.extract_keywords_from_text(text)
    keywords = rake_nltk_var.get_ranked_phrases()[:5]

    return summarize_text, keywords


def get_txt(file_path):
    with open(file_path, 'r') as txtf:
        text = txtf.read().replace("\n", " ")
        return text


def get_pdf(file_path):
    output_string = StringIO()
    with open(file_path, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(
            rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)

        return (output_string.getvalue())


def get_readability(intxt):

    letterno = float(letters(intxt))
    wordsno = float(words(intxt))
    sentenceno = float(sentences(intxt))

    avgl = float((letterno / wordsno)) * 100
    s = float((sentenceno / wordsno)) * 100

    # https://en.wikipedia.org/wiki/Coleman%E2%80%93Liau_index

    index = int(round((0.0588 * avgl) - (0.296 * s) - 15.8))

    if index < 1:
        return "Readability Index: Before Grade 1"
    elif index >= 13 and index < 16:
        return "Readability Index: College Level"
    elif index >= 16:
        return "Readability Index: Expert"
    else:
        return f"Readability Index: Grade {index}"


def letters(string):

    lettercount = 0
    string = string.upper()

    for char in string:
        # using ord to convert from str to int of ascii
        word_ascii = ord(char) - 65
        if word_ascii < 0 or word_ascii > 25:
            pass
        else:
            lettercount += 1
    return lettercount


def words(string):
    wordcount = 1

    for char in string:
        # using ord to convert from str to int of ascii
        char_ascii = ord(char)

        if char_ascii == 32:
            wordcount += 1

    return wordcount


def sentences(string):
    sentencecount = 0

    for char in string:
        # using ord to convert from str to int of ascii
        char = ord(char)

        if char == 33 or char == 63 or char == 46:
            sentencecount += 1

    return sentencecount


def check_plagrism(txt1, txt2):
    return f"The texts are {int(SequenceMatcher(None, txt1, txt2).ratio() * 100)}% similar"


def get_docx(file_path):
    extracted_txt = ""
    doc = docx.Document(file_path)
    for para in doc.paragraphs:
        extracted_txt += str(para.text)
    return extracted_txt


blacklist = ['[document]', 'noscript', 'header',
             'html', 'meta', 'head', 'input', 'script']


def get_epub(file_path):
    txt = str(epub2text(file_path)[1].replace("\n", "").rstrip())
    return txt


def epub2thtml(epub_path):
    book = epub.read_epub(epub_path)
    chapters = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            chapters.append(item.get_content())
    return chapters


def chap2text(chap):
    output = ''
    soup = BeautifulSoup(chap, 'html.parser')
    text = soup.find_all(text=True)
    for t in text:
        if t.parent.name not in blacklist:
            output += '{} '.format(t)
    return output


def thtml2ttext(thtml):
    Output = []
    for html in thtml:
        text = chap2text(html)
        Output.append(text)
    return Output


def epub2text(epub_path):
    chapters = epub2thtml(epub_path)
    ttext = thtml2ttext(chapters)
    return ttext


# from youtube_transcript_api import YouTubeTranscriptApi

# def get_captions(lnk):
#     if "youtube.com/watch?v" in lnk:
#         lnk = lnk.split("https://www.youtube.com/watch?v=")[1]
#     elif "https://youtu.be/" in lnk:
#         lnk = lnk.split("https://youtu.be/")[1]
#     else:
#         return "Invalid link"

#     try:
#         srt = YouTubeTranscriptApi.get_transcript(lnk,
#                                             languages=['en'])
#     except:
#         return "Couldn't find caption"
#     captions = ""
#     for i in srt:
#         captions += i["text"] + " "

#     return captions
